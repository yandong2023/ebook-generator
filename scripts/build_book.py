#!/usr/bin/env python3
"""Build EPUB3, DOCX, HTML and optional PDF from an ebook-generator project."""

from __future__ import annotations

import argparse
import base64
import html
import mimetypes
import re
import sys
import uuid
import zipfile
from datetime import date, datetime, timezone
from pathlib import Path
from urllib.parse import urlparse

import mistune
import yaml
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches

MD = mistune.create_markdown(escape=False, plugins=["table", "footnotes", "strikethrough"])


def config(project: Path) -> dict:
    path = project / "book.yaml"
    if not path.exists():
        raise FileNotFoundError(f"Missing {path}")
    data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    defaults = {
        "title": "Untitled Ebook", "subtitle": "", "author": "Unknown Author",
        "language": "en", "publisher": "", "description": "", "identifier": "",
        "date": "", "cover": "assets/cover.jpg", "output_name": "book",
    }
    return defaults | data


def chapters(project: Path) -> list[Path]:
    files = sorted((project / "chapters").glob("*.md"))
    if not files:
        raise FileNotFoundError("No Markdown chapters found under chapters/")
    return files


def safe_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip()).strip("-._") or "book"


def media_type(path: Path) -> str:
    return mimetypes.guess_type(path.name)[0] or "application/octet-stream"


def local_image(src: str, chapter: Path, project: Path) -> Path | None:
    if urlparse(src).scheme in {"http", "https", "data"} or src.startswith("#"):
        return None
    src = src.split("#", 1)[0].split("?", 1)[0]
    for candidate in (chapter.parent / src, project / src):
        candidate = candidate.resolve()
        if candidate.exists():
            return candidate
    return (project / src).resolve()


def title_from(fragment: str, fallback: str) -> str:
    h1 = BeautifulSoup(fragment, "html.parser").find("h1")
    return h1.get_text(" ", strip=True) if h1 else fallback


def prepare_fragment(fragment: str, chapter: Path, project: Path, assets: dict[Path, str]) -> str:
    soup = BeautifulSoup(fragment, "html.parser")
    for img in list(soup.find_all("img")):
        src = img.get("src")
        if not src:
            continue
        local = local_image(src, chapter, project)
        if local is None:
            p = soup.new_tag("p")
            p.string = f"[Remote image omitted: {img.get('alt') or src}]"
            img.replace_with(p)
            continue
        if not local.exists():
            raise FileNotFoundError(f"Missing image in {chapter.name}: {src}")
        if local not in assets:
            assets[local] = f"images/{len(assets)+1:03d}-{safe_name(local.name)}"
        img["src"] = "../" + assets[local]
    return soup.decode(formatter="minimal")


def make_epub(project: Path, cfg: dict, files: list[Path], out: Path, css: str) -> None:
    assets: dict[Path, str] = {}
    docs: list[tuple[str, str, str]] = []
    for i, chapter in enumerate(files, 1):
        raw = MD(chapter.read_text(encoding="utf-8"))
        title = title_from(raw, chapter.stem)
        fragment = prepare_fragment(raw, chapter, project, assets)
        href = f"chapters/{i:02d}-{safe_name(chapter.stem)}.xhtml"
        xhtml = f'''<?xml version="1.0" encoding="UTF-8"?>\n<!DOCTYPE html>\n<html xmlns="http://www.w3.org/1999/xhtml" lang="{html.escape(str(cfg['language']))}"><head><meta charset="utf-8"/><title>{html.escape(title)}</title><link rel="stylesheet" href="../styles/ebook.css" type="text/css"/></head><body>{fragment}</body></html>'''
        docs.append((href, title, xhtml))

    cover_path = (project / str(cfg.get("cover") or "")).resolve() if cfg.get("cover") else None
    cover_href = None
    if cover_path and cover_path.exists():
        cover_href = "images/cover" + (cover_path.suffix.lower() or ".jpg")
        assets[cover_path] = cover_href

    identifier = str(cfg.get("identifier") or "").strip() or "urn:uuid:" + str(
        uuid.uuid5(uuid.NAMESPACE_URL, f"ebook-generator:{cfg['title']}:{cfg['author']}")
    )
    modified = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    pub_date = str(cfg.get("date") or date.today().isoformat())

    manifest = [
        '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>',
        '<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>',
        '<item id="css" href="styles/ebook.css" media-type="text/css"/>',
    ]
    spine: list[str] = []
    if cover_href and cover_path:
        manifest += [
            f'<item id="cover-image" href="{cover_href}" media-type="{media_type(cover_path)}" properties="cover-image"/>',
            '<item id="cover-page" href="cover.xhtml" media-type="application/xhtml+xml"/>',
        ]
        spine.append('<itemref idref="cover-page"/>')
    for i, (href, _title, _doc) in enumerate(docs, 1):
        manifest.append(f'<item id="c{i}" href="{href}" media-type="application/xhtml+xml"/>')
        spine.append(f'<itemref idref="c{i}"/>')
    for path, href in assets.items():
        if cover_path and path == cover_path:
            continue
        manifest.append(f'<item id="img{len(manifest)}" href="{href}" media-type="{media_type(path)}"/>')

    extra_meta = ""
    if cfg.get("publisher"):
        extra_meta += f'<dc:publisher>{html.escape(str(cfg["publisher"]))}</dc:publisher>'
    if cfg.get("description"):
        extra_meta += f'<dc:description>{html.escape(str(cfg["description"]))}</dc:description>'
    opf = f'''<?xml version="1.0" encoding="UTF-8"?>\n<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="book-id"><metadata xmlns:dc="http://purl.org/dc/elements/1.1/"><dc:identifier id="book-id">{html.escape(identifier)}</dc:identifier><dc:title>{html.escape(str(cfg['title']))}</dc:title><dc:creator>{html.escape(str(cfg['author']))}</dc:creator><dc:language>{html.escape(str(cfg['language']))}</dc:language><dc:date>{html.escape(pub_date)}</dc:date><meta property="dcterms:modified">{modified}</meta>{extra_meta}</metadata><manifest>{''.join(manifest)}</manifest><spine toc="ncx">{''.join(spine)}</spine></package>'''
    nav_items = "".join(f'<li><a href="{href}">{html.escape(title)}</a></li>' for href, title, _ in docs)
    nav = f'''<?xml version="1.0" encoding="UTF-8"?>\n<!DOCTYPE html>\n<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops"><head><title>Contents</title></head><body><nav epub:type="toc" id="toc"><h1>Contents</h1><ol>{nav_items}</ol></nav></body></html>'''
    points = "".join(f'<navPoint id="n{i}" playOrder="{i}"><navLabel><text>{html.escape(title)}</text></navLabel><content src="{href}"/></navPoint>' for i, (href, title, _) in enumerate(docs, 1))
    ncx = f'''<?xml version="1.0" encoding="UTF-8"?>\n<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1"><head><meta name="dtb:uid" content="{html.escape(identifier)}"/></head><docTitle><text>{html.escape(str(cfg['title']))}</text></docTitle><navMap>{points}</navMap></ncx>'''
    container = '<?xml version="1.0"?><container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container"><rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles></container>'

    out.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(out, "w") as zf:
        zf.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        for name, data in {
            "META-INF/container.xml": container,
            "OEBPS/content.opf": opf,
            "OEBPS/nav.xhtml": nav,
            "OEBPS/toc.ncx": ncx,
            "OEBPS/styles/ebook.css": css,
        }.items():
            zf.writestr(name, data, compress_type=zipfile.ZIP_DEFLATED)
        if cover_href and cover_path:
            cover = f'''<?xml version="1.0" encoding="UTF-8"?><html xmlns="http://www.w3.org/1999/xhtml"><head><title>Cover</title></head><body><div><img alt="Cover" src="{cover_href}"/></div></body></html>'''
            zf.writestr("OEBPS/cover.xhtml", cover, compress_type=zipfile.ZIP_DEFLATED)
        for href, _title, doc in docs:
            zf.writestr("OEBPS/" + href, doc, compress_type=zipfile.ZIP_DEFLATED)
        for path, href in assets.items():
            zf.writestr("OEBPS/" + href, path.read_bytes(), compress_type=zipfile.ZIP_DEFLATED)


def html_with_embedded_images(fragment: str, chapter: Path, project: Path) -> str:
    soup = BeautifulSoup(fragment, "html.parser")
    for img in soup.find_all("img"):
        src = img.get("src")
        if not src:
            continue
        local = local_image(src, chapter, project)
        if local and local.exists():
            encoded = base64.b64encode(local.read_bytes()).decode("ascii")
            img["src"] = f"data:{media_type(local)};base64,{encoded}"
    return str(soup)


def make_html(project: Path, cfg: dict, files: list[Path], out: Path, css: str) -> None:
    body = []
    for chapter in files:
        fragment = MD(chapter.read_text(encoding="utf-8"))
        body.append('<section class="chapter">' + html_with_embedded_images(fragment, chapter, project) + '</section>')
    doc = f'''<!doctype html><html lang="{html.escape(str(cfg['language']))}"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>{html.escape(str(cfg['title']))}</title><style>{css}.chapter{{page-break-before:always}}.chapter:first-of-type{{page-break-before:auto}}</style></head><body>{''.join(body)}</body></html>'''
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(doc, encoding="utf-8")


def add_runs(paragraph, node: Tag | NavigableString) -> None:
    if isinstance(node, NavigableString):
        paragraph.add_run(str(node))
        return
    if not isinstance(node, Tag):
        return
    before = len(paragraph.runs)
    for child in node.children:
        add_runs(paragraph, child)
    after = len(paragraph.runs)
    for run in paragraph.runs[before:after]:
        if node.name in {"strong", "b"}: run.bold = True
        if node.name in {"em", "i"}: run.italic = True
        if node.name == "code": run.font.name = "Courier New"


def make_docx(project: Path, cfg: dict, files: list[Path], out: Path) -> None:
    doc = Document()
    doc.core_properties.title = str(cfg["title"])
    doc.core_properties.author = str(cfg["author"])
    doc.add_heading(str(cfg["title"]), level=0)
    if cfg.get("subtitle"): doc.add_paragraph(str(cfg["subtitle"]))
    doc.add_paragraph(str(cfg["author"]))
    doc.add_page_break()
    for n, chapter in enumerate(files):
        soup = BeautifulSoup(MD(chapter.read_text(encoding="utf-8")), "html.parser")
        for node in soup.contents:
            if not isinstance(node, Tag):
                continue
            if node.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
                doc.add_heading(node.get_text(" ", strip=True), level=min(int(node.name[1]), 6))
            elif node.name == "p":
                p = doc.add_paragraph(); add_runs(p, node)
                for img in node.find_all("img"):
                    local = local_image(img.get("src", ""), chapter, project)
                    if local and local.exists():
                        try: doc.add_picture(str(local), width=Inches(5.8))
                        except Exception: pass
            elif node.name in {"ul", "ol"}:
                style = "List Bullet" if node.name == "ul" else "List Number"
                for li in node.find_all("li", recursive=False):
                    p = doc.add_paragraph(style=style); add_runs(p, li)
            elif node.name == "pre":
                p = doc.add_paragraph(); r = p.add_run(node.get_text()); r.font.name = "Courier New"
            elif node.name == "table":
                rows = node.find_all("tr"); width = max((len(r.find_all(["th", "td"], recursive=False)) for r in rows), default=0)
                if rows and width:
                    table = doc.add_table(rows=len(rows), cols=width); table.style = "Table Grid"
                    for ri, row in enumerate(rows):
                        for ci, cell in enumerate(row.find_all(["th", "td"], recursive=False)):
                            table.cell(ri, ci).text = cell.get_text(" ", strip=True)
        if n < len(files)-1: doc.add_page_break()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def make_pdf(html_path: Path, out: Path) -> None:
    try:
        from weasyprint import HTML
    except ImportError as exc:
        raise RuntimeError("PDF export requires: pip install -r requirements-pdf.txt") from exc
    HTML(filename=str(html_path), base_url=str(html_path.parent)).write_pdf(str(out))


def main() -> int:
    parser = argparse.ArgumentParser(description="Build an ebook-generator project")
    parser.add_argument("project")
    parser.add_argument("--format", choices=["epub", "docx", "html", "pdf", "all"], default="all")
    parser.add_argument("--output-name")
    args = parser.parse_args()

    project = Path(args.project).expanduser().resolve()
    cfg = config(project); files = chapters(project)
    css = (Path(__file__).resolve().parents[1] / "assets" / "epub.css").read_text(encoding="utf-8")
    name = safe_name(args.output_name or str(cfg.get("output_name") or cfg["title"]))
    dist = project / "dist"; dist.mkdir(parents=True, exist_ok=True)
    wanted = [args.format] if args.format != "all" else ["epub", "docx", "html", "pdf"]
    html_path = dist / f"{name}.html"

    for fmt in wanted:
        try:
            if fmt == "epub": out = dist / f"{name}.epub"; make_epub(project, cfg, files, out, css)
            elif fmt == "docx": out = dist / f"{name}.docx"; make_docx(project, cfg, files, out)
            elif fmt == "html": out = html_path; make_html(project, cfg, files, out, css)
            else:
                if not html_path.exists(): make_html(project, cfg, files, html_path, css)
                out = dist / f"{name}.pdf"; make_pdf(html_path, out)
            print(f"Built {fmt.upper()}: {out}")
        except RuntimeError as exc:
            if fmt == "pdf" and args.format == "all":
                print(f"Skipped PDF: {exc}", file=sys.stderr)
                continue
            raise
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
